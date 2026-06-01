from docx import Document
from docx.text.paragraph import Paragraph # For type checking
from docx.table import Table, _Cell     # For type checking
import os
import html # Import the html module for unescaping

class WordHandler:
    def __init__(self, translator):
        self.translator = translator

    def set_font(self, element, font_name):
        """
        Sets the font name for text within the given element.
        This method is adapted from PptHandler's set_font, focusing on docx structures.
        In docx, font properties are typically set on 'runs'.
        """
        if isinstance(element, Paragraph):
            for run in element.runs:
                run.font.name = font_name
        elif isinstance(element, _Cell):
            # A cell contains paragraphs, so iterate through them
            for paragraph in element.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
        else:
            # For other elements, print a warning or handle as appropriate.
            # The original PptHandler `set_font` had a `pass` for GraphicFrame.
            print(f"Warning: set_font not directly applicable or implemented for element type {type(element)}")


    def _translate_text_in_run(self, run):
        """
        Translates the text content of a docx Run object.
        """
        if run.text and str(run.text).strip():
            to_translate = str(run.text).strip()
            translated = self.translator.translate(to_translate)
            # Unescape HTML entities if the translator returns them
            if translated:
                translated = html.unescape(translated)
            print(f"\tRun text: {to_translate!r} -> {translated!r}")
            if translated:
                run.text = translated

    def translate(self, path):
        """
        Translates the text content of a .docx document.
        """
        if not path.endswith(".docx"):
            print("Warning: The provided path does not end with '.docx'. Attempting to open anyway.")

        document = Document(path)
        print("Translating document content...")

        # Process main document body
        print("Processing main document body...")
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                self._translate_text_in_run(run)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._translate_text_in_run(run)

        # Process headers and footers in all sections
        print("Processing headers and footers...")
        for section in document.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for run in paragraph.runs:
                    self._translate_text_in_run(run)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self._translate_text_in_run(run)

            # Footer
            for paragraph in section.footer.paragraphs:
                for run in paragraph.runs:
                    self._translate_text_in_run(run)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self._translate_text_in_run(run)

        # Determine the output path, similar to the PptHandler logic
        if path.endswith(".docx"):
            base, ext = os.path.splitext(path)
            output_path = f"{base}-tr{ext}"
        else:
            output_path = f"{path}-tr.docx" # Default to adding .docx if original didn't have it

        document.save(output_path)
        print(f"Translated document saved to {output_path}")